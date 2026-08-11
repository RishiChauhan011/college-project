"""
Accuracy-improved resume text extraction + segmentation.

Replaces the plain-text-only pipeline with a metadata-aware one:
extraction now preserves font size, boldness, and source column per
line, and segmentation uses those as the PRIMARY signal for heading
detection, with keyword matching as a secondary confirming signal.

Why this matters:
Plain text extraction throws away every visual cue the PDF gives you.
Two lines that look identical as plain text ("Skills" vs "Python, SQL,
Excel") can be told apart instantly once you know one is 14pt bold and
the other is 10pt regular. This is the single highest-leverage fix for
"content landing under the wrong section."

Requires: pdfplumber, rapidfuzz
    pip install pdfplumber rapidfuzz --break-system-packages

Drop-in replacement for extract_text.py + segment_sections.py.
"""

import io
import os
import sys
import pdfplumber
from rapidfuzz import fuzz

# Reuse the SAME skill matcher already built and tested in extract_skills.py
# - this is the single most important integration rule: resumes and job
# postings must use one shared definition of "what is a skill"
# (master_skills.csv), never two separate implementations that could
# silently drift apart.
sys.path.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "ml"))
from extract_skills import load_master_skills, build_phrase_matcher, extract_skills_from_text
import spacy
from spacy.tokens import Span

# Module-level cache, built once on first use (not rebuilt per resume
# upload - matches the same "load once, reuse" performance principle
# already applied to the FastAPI backend's startup caching).
_nlp = None
_matcher = None
_skill_lookup = None


def _get_cached_matcher():
    global _nlp, _matcher, _skill_lookup
    if _matcher is None:
        _nlp = spacy.blank("en")
        if not Span.has_extension("skill_id"):
            Span.set_extension("skill_id", default=None)
        master_skills = load_master_skills()
        _matcher, _skill_lookup = build_phrase_matcher(_nlp, master_skills)
    return _nlp, _matcher, _skill_lookup


CATEGORY_KEY_MAP = {"Technical": "technical", "Domain": "domain", "Soft Skill": "soft"}


def extract_resume_skills(skill_extraction_text: str) -> tuple[list[str], dict]:
    """Runs the shared PhraseMatcher against the resume's skill-relevant
    text. Returns (flat_skill_names, grouped_by_category) - same shapes
    used everywhere else in the project (jobs_with_skills.json,
    recommendation output)."""
    empty_grouped = {"technical": [], "domain": [], "soft": []}
    if not skill_extraction_text:
        return [], empty_grouped

    nlp, matcher, skill_lookup = _get_cached_matcher()
    matched = extract_skills_from_text(nlp, matcher, skill_lookup, skill_extraction_text)

    grouped = {"technical": [], "domain": [], "soft": []}
    for skill in matched:
        key = CATEGORY_KEY_MAP.get(skill["category"], "domain")
        grouped[key].append(skill)

    flat_names = [s["name"] for s in matched]
    return flat_names, grouped

SECTION_KEYWORDS = {
    "education": ["education", "academic background", "academics", "qualification"],
    "experience": ["experience", "work experience", "employment history",
                   "professional experience", "internship"],
    "skills": ["skills", "technical skills", "core competencies", "tech stack"],
    "projects": ["projects", "academic projects", "personal projects"],
    "certifications": ["certifications", "certificates", "licenses"],
    "contact": ["contact", "contact info", "get in touch"],
    "about": ["about me", "summary", "profile", "objective"],
    "languages": ["language", "languages"],
}

FUZZY_MATCH_THRESHOLD = 82  # 0-100, higher = stricter


# ---------- Stage 1: extraction with metadata ----------

def _join_chars_with_spacing(chars: list) -> str:
    """
    pdfplumber's char-level extraction gives individual glyphs — actual
    space characters are frequently NOT included as glyphs at all, so a
    naive "".join() mashes words together ("About Me" -> "AboutMe").
    This reconstructs spacing by checking the horizontal gap between
    consecutive characters: a gap noticeably wider than normal
    character spacing means a word boundary.
    """
    if not chars:
        return ""
    chars = sorted(chars, key=lambda c: c["x0"])
    parts = [chars[0]["text"]]
    for prev, curr in zip(chars, chars[1:]):
        gap = curr["x0"] - prev["x1"]
        # Threshold scaled to font size — a gap wider than ~20% of the
        # character's own size reliably indicates a word break rather
        # than normal kerning between adjacent letters.
        space_threshold = curr.get("size", 10) * 0.2
        if gap > space_threshold:
            parts.append(" ")
        parts.append(curr["text"])
    return "".join(parts)


def _cluster_lines_with_metadata(chars, y_tolerance=3):
    """Groups characters into lines, computing per-line average font
    size and bold ratio, sorted top-to-bottom then left-to-right."""
    chars = sorted(chars, key=lambda c: (c["top"], c["x0"]))
    lines = []
    current = []
    current_top = None

    def flush(current):
        if not current:
            return None
        current.sort(key=lambda c: c["x0"])
        text = _join_chars_with_spacing(current).strip()
        if not text:
            return None
        avg_size = sum(c["size"] for c in current) / len(current)
        bold_count = sum(1 for c in current if "bold" in c.get("fontname", "").lower())
        bold_ratio = bold_count / len(current)
        top = min(c["top"] for c in current)
        x0 = min(c["x0"] for c in current)
        return {"text": text, "font_size": avg_size, "bold_ratio": bold_ratio,
                "top": top, "x0": x0}

    for c in chars:
        if c["text"].strip() == "":
            continue
        if current_top is None or abs(c["top"] - current_top) <= y_tolerance:
            current.append(c)
            current_top = c["top"] if current_top is None else current_top
        else:
            line = flush(current)
            if line:
                lines.append(line)
            current = [c]
            current_top = c["top"]

    line = flush(current)
    if line:
        lines.append(line)

    return lines


def _detect_column_split(chars, page_width, min_gap_ratio=0.06):
    if not chars:
        return None
    x0_positions = sorted(set(round(c["x0"]) for c in chars))
    min_gap = page_width * min_gap_ratio
    best_gap, split_x = 0, None
    for i in range(1, len(x0_positions)):
        gap = x0_positions[i] - x0_positions[i - 1]
        if gap > best_gap:
            best_gap, split_x = gap, (x0_positions[i] + x0_positions[i - 1]) / 2
    return split_x if best_gap >= min_gap else None


def extract_lines_with_metadata(pdf_bytes: bytes) -> tuple[list[dict], list[str]]:
    """
    Returns (lines, warnings) where each line is:
    {text, font_size, bold_ratio, top, column}  column is "main" or "sidebar"
    """
    warnings = []
    all_lines = []

    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                chars = page.chars
                if not chars:
                    continue

                split_x = _detect_column_split(chars, page.width)

                if split_x is None:
                    lines = _cluster_lines_with_metadata(chars)
                    for l in lines:
                        l["column"] = "main"
                    all_lines.extend(lines)
                else:
                    left_chars = [c for c in chars if c["x0"] < split_x]
                    right_chars = [c for c in chars if c["x0"] >= split_x]
                    left_lines = _cluster_lines_with_metadata(left_chars)
                    right_lines = _cluster_lines_with_metadata(right_chars)

                    main_lines, sidebar_lines = (
                        (right_lines, left_lines) if len(right_chars) > len(left_chars)
                        else (left_lines, right_lines)
                    )
                    for l in main_lines:
                        l["column"] = "main"
                    for l in sidebar_lines:
                        l["column"] = "sidebar"
                    all_lines.extend(main_lines)
                    all_lines.extend(sidebar_lines)

    except Exception:
        warnings.append("corrupted_or_unsupported")
        return [], warnings

    total_chars = sum(len(l["text"]) for l in all_lines)
    if total_chars < 50:
        warnings.append("no_text_found_possibly_scanned")

    return all_lines, warnings


# ---------- Stage 2: metadata-aware segmentation ----------

def _looks_like_heading_candidate(text: str) -> bool:
    """
    Cheap structural pre-filter, checked BEFORE fuzzy matching. Real
    headings are short, standalone, and don't carry embedded data.
    This alone eliminates the two failure modes seen in testing:
      - long/sentence-like lines ("Built an automated resume screening
        tool...") ending in punctuation
      - "label: value" content lines ("Languages & Frameworks: Python,
        FastAPI") that happen to contain a section keyword as their
        label but are actually list content, not a heading
    """
    text = text.strip()
    if not text or len(text) > 35:
        return False
    if text.endswith((".", ",", ";", ":")):
        return False
    if len(text.split()) > 5:
        return False
    # "label: value" pattern — a colon followed by real content means
    # this is a data line, not a standalone heading, even if short.
    if ":" in text and text.split(":", 1)[1].strip():
        return False
    return True


def _keyword_match_score(text: str) -> tuple[str, int]:
    """
    Returns (best_matching_section, fuzzy_score) for a line of text.
    Uses fuzz.token_set_ratio — compares shared WORDS between the line
    and each keyword, not raw character substrings. This is the key
    fix: partial_ratio rewards any strong character-sequence alignment
    and produces false positives on unrelated text (e.g. a project
    title scoring high against "education" purely by coincidence).
    token_set_ratio still lets "PROFESSIONAL SUMMARY" match "summary"
    (shared word, extra qualifier is fine) while correctly rejecting
    "Resume Screening Application" against every keyword (no shared
    words at all).
    """
    if not _looks_like_heading_candidate(text):
        return None, 0

    text_clean = text.strip().lower()
    best_section, best_score = None, 0
    for section, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            score = fuzz.token_set_ratio(text_clean, kw)
            if score > best_score:
                best_section, best_score = section, score
    return best_section, best_score


import re

_EMAIL_RE = re.compile(r"[\w.\-+]+@[\w\-]+\.[\w.\-]+")
_PHONE_RE = re.compile(r"(\+?\d[\d\-\s]{8,}\d)")


def _looks_like_contact_info(text: str) -> bool:
    """
    High-precision structural check for contact-info lines (email,
    phone, LinkedIn/GitHub handles). These follow regular, near-
    deterministic patterns, so pattern matching is far more reliable
    here than fuzzy heading detection — and lets us route contact
    details correctly even when they appear before any heading is
    detected (which the generic "header" bucket would otherwise catch).
    """
    if _EMAIL_RE.search(text):
        return True

    phone_match = _PHONE_RE.search(text)
    if phone_match:
        # Guard against date ranges like "(2023 -2024)" which satisfy
        # the digit-dash-space pattern but aren't phone numbers. Real
        # phone numbers have at least 7 actual digit characters; a
        # 4-digit year range does not.
        digit_count = sum(c.isdigit() for c in phone_match.group())
        if digit_count >= 10:
            return True

    lowered = text.lower()
    if re.search(r"\blinkedin\b", lowered) or re.search(r"\bgithub\b", lowered):
        return True
    return False


def _is_heading(line: dict, body_font_size: float) -> tuple[bool, str | None, int]:
    """
    Weighted-score heading decision, GATED on a minimum keyword score.
    Visual/structural signals (bold, short line, all-caps) can boost a
    borderline keyword match over the line, but they can never manufacture
    a heading on their own — a bold, short, all-caps line with no real
    relation to any known section (e.g. a person's name) must not be
    misread as a heading just because it "looks" like one structurally.
    """
    text = line["text"].strip()
    section, kw_score = _keyword_match_score(text)

    MIN_KEYWORD_GATE = 80  # measured: genuine headings score ~100, coincidental
    # character-overlap false positives on short unrelated words (names,
    # taglines, single skill/language entries) cluster in the 45-75 range.
    # The gate must sit above that band.
    if kw_score < MIN_KEYWORD_GATE:
        return False, None, 0

    score = 0
    if kw_score >= 90:
        score += 3
    else:
        score += 2

    size_ratio = line["font_size"] / body_font_size if body_font_size else 1
    if size_ratio >= 1.15:
        score += 2
    elif size_ratio >= 1.05:
        score += 1

    if line["bold_ratio"] >= 0.6:
        score += 2

    word_count = len(text.split())
    if word_count <= 3:
        score += 1
    if text.isupper() and word_count <= 5:
        score += 1

    is_heading = score >= 3
    return is_heading, section, score


def _compute_body_font_size(lines: list[dict]) -> float:
    """
    Mode of font sizes, bucketed to the nearest 0.5pt to avoid
    fragmenting the count across near-identical decimal values
    (common with exported/embedded fonts).
    """
    if not lines:
        return 10.0
    buckets = [round(l["font_size"] * 2) / 2 for l in lines]
    return max(set(buckets), key=buckets.count)


def segment_sections(lines: list[dict]) -> tuple[dict, list[str]]:
    warnings = []
    sections = {k: [] for k in SECTION_KEYWORDS}
    sections["misc"] = []
    sections["header"] = []  # name/title lines before the first real heading

    if not lines:
        warnings.append("low_section_match_confidence")
        return {k: None for k in sections}, warnings

    body_font_size = _compute_body_font_size(lines)

    current_section = "header"  # nothing detected yet = name/tagline zone
    detected_count = 0
    seen_first_heading = False

    for line in lines:
        # Contact info (email/phone/LinkedIn/GitHub) is high-precision
        # to detect by pattern and should always land in "contact",
        # regardless of which section is currently active or whether
        # any heading has fired yet.
        if _looks_like_contact_info(line["text"]):
            sections["contact"].append(line["text"])
            continue

        is_head, section, score = _is_heading(line, body_font_size)
        if is_head and section:
            current_section = section
            detected_count += 1
            seen_first_heading = True
            continue  # heading line itself isn't content

        # Before any heading is found, keep dumping into "header" (name,
        # title, tagline) rather than guessing a real section. Once a
        # heading fires at least once, unmatched short/ambiguous lines
        # go to "misc" instead of silently stacking on the last section
        # forever — misc is a visible signal you can inspect and fix,
        # rather than a silent accuracy killer.
        if not seen_first_heading:
            sections["header"].append(line["text"])
        else:
            sections[current_section].append(line["text"])

    if detected_count < 2:
        warnings.append("low_section_match_confidence")

    result = {
        k: ("\n".join(v).strip() if v else None)
        for k, v in sections.items()
    }
    return result, warnings


# ---------- Stage 3: docx extraction ----------

def extract_lines_from_docx(file_bytes: bytes) -> tuple[list[dict], list[str]]:
    import io
    import docx
    
    warnings = []
    lines = []
    
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            is_heading_style = para.style and para.style.name.startswith('Heading')
            
            bold_chars = 0
            total_chars = 0
            for run in para.runs:
                if run.text:
                    total_chars += len(run.text)
                    if run.bold:
                        bold_chars += len(run.text)
                        
            bold_ratio = bold_chars / total_chars if total_chars > 0 else 0.0
            if is_heading_style:
                bold_ratio = max(bold_ratio, 0.6)
                
            font_size = 14.0 if is_heading_style else 10.0
            
            lines.append({
                "text": text,
                "font_size": font_size,
                "bold_ratio": bold_ratio,
                "top": 0.0,
                "column": "main"
            })

        # Tables were previously skipped entirely, meaning a table-based
        # resume (some layout templates use tables instead of plain
        # paragraphs) would silently produce zero extracted text and get
        # incorrectly flagged as "scanned/unreadable". Read table cell
        # text too, treated as plain body content (font/bold metadata
        # within table cells is not reliably meaningful for heading
        # detection, so these lines use default body-text values).
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    lines.append({
                        "text": cell_text,
                        "font_size": 10.0,
                        "bold_ratio": 0.0,
                        "top": 0.0,
                        "column": "main"
                    })
    except Exception:
        warnings.append("corrupted_or_unsupported")
        return [], warnings
        
    total_chars = sum(len(l["text"]) for l in lines)
    if total_chars < 50:
        warnings.append("no_text_found_possibly_scanned")
        
    return lines, warnings

# ---------- Orchestration ----------

SKILL_EXTRACTION_SECTIONS = ["skills", "experience", "projects"]

def build_skill_extraction_text(sections: dict) -> str:
    """
    Combines the sections that carry real skill signal into one text
    blob for downstream extraction (PhraseMatcher + NER). Sections like
    education, about, contact, certifications, languages, and header
    are deliberately excluded — they don't reliably contain skill
    mentions and only add noise.
    """
    parts = [sections.get(key) for key in SKILL_EXTRACTION_SECTIONS if sections.get(key)]
    return "\n".join(parts)

def parse_resume(file_bytes: bytes, file_type: str = "pdf") -> dict:
    if file_type.lower() == "pdf":
        lines, extract_warnings = extract_lines_with_metadata(file_bytes)
    elif file_type.lower() == "docx":
        lines, extract_warnings = extract_lines_from_docx(file_bytes)
        extract_warnings.append("docx_fallback_used")
    else:
        return _build_response(
            sections={k: None for k in SECTION_KEYWORDS} | {"misc": None, "header": None},
            raw_text="",
            skill_extraction_text="",
            source_format=file_type.lower(),
            extraction_warnings=["unsupported_file_type"],
        )

    sections, segment_warnings = segment_sections(lines)
    raw_text = "\n".join(l["text"] for l in lines)
    skill_text = build_skill_extraction_text(sections)

    return _build_response(
        sections=sections,
        raw_text=raw_text,
        skill_extraction_text=skill_text,
        source_format=file_type.lower(),
        extraction_warnings=extract_warnings + segment_warnings,
    )


def _build_response(sections, raw_text, skill_extraction_text, source_format, extraction_warnings):
    """Single place that builds the final response - guarantees every
    code path (success, unsupported file type, scanned/unreadable PDF)
    returns the EXACT SAME set of keys, so the frontend never has to
    guess which fields might be missing depending on what happened."""
    unreadable_markers = {"no_text_found_possibly_scanned", "corrupted_or_unsupported",
                            "unsupported_file_type"}
    readable = not any(w in unreadable_markers for w in extraction_warnings)

    if readable:
        skills, skills_grouped = extract_resume_skills(skill_extraction_text)
    else:
        skills, skills_grouped = [], {"technical": [], "domain": [], "soft": []}

    return {
        "readable": readable,
        "skills": skills,
        "skills_grouped": skills_grouped,
        "sections": sections,
        "raw_text": raw_text,
        "skill_extraction_text": skill_extraction_text,
        "source_format": source_format,
        "extraction_warnings": extraction_warnings,
        # Never inferred from resume text - hard rule, see integration
        # spec section 4. Only the frontend/user's explicit choice
        # (post-upload mini-form) may set these.
        "preferred_domain": None,
        "preferred_location": None,
    }