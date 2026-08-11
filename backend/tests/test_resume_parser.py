"""
test_resume_parser.py

Tests the ACTUAL current parse_resume.py implementation (consolidated
single-file parser with metadata-aware extraction + segmentation +
integrated skill extraction via the shared PhraseMatcher).

The previous version of this file imported from resume_parser.extract_text
and resume_parser.segment_sections as separate modules with different
function names - those no longer exist; the parser was consolidated into
one file with a different, metadata-based API. This version tests the
real, current public interface.

Requires: pip install pytest reportlab python-docx
"""

import io
import pytest
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import docx

from resume_parser.parse_resume import parse_resume, extract_resume_skills


def create_synthetic_pdf(text_content):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    y = 750
    for line in text_content.split('\n'):
        if line.strip():
            c.drawString(72, y, line)
        y -= 15
        if y < 72:
            c.showPage()
            y = 750
    c.save()
    return buffer.getvalue()


def create_synthetic_docx(text_content, use_tables=False):
    document = docx.Document()
    if use_tables:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Content"
        hdr_cells[1].text = ""
        for line in text_content.split('\n'):
            if line.strip():
                row_cells = table.add_row().cells
                row_cells[0].text = line
                row_cells[1].text = ""
    else:
        for line in text_content.split('\n'):
            if line.strip():
                document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


SAMPLE_TEXT_1 = """John Doe
john.doe@email.com
555-0100

EXPERIENCE

Software Engineer
Worked on backend systems.

EDUCATION

B.S. Computer Science
University of Tech

SKILLS

Python, SQL, Docker, Machine Learning
"""


# ---------------------------------------------------------------------------
# Contract tests: EVERY response, regardless of path taken, must have
# exactly this shape (this is the normalized-output-contract requirement)
# ---------------------------------------------------------------------------

REQUIRED_KEYS = {
    "readable", "skills", "skills_grouped", "sections", "raw_text",
    "skill_extraction_text", "source_format", "extraction_warnings",
    "preferred_domain", "preferred_location",
}


def assert_valid_contract(result: dict):
    """Every parse_resume() response must have this exact set of keys,
    regardless of whether the file was readable, unsupported, or scanned."""
    assert set(result.keys()) == REQUIRED_KEYS, (
        f"Response keys don't match the normalized contract. "
        f"Missing: {REQUIRED_KEYS - set(result.keys())}, "
        f"Unexpected: {set(result.keys()) - REQUIRED_KEYS}"
    )
    assert isinstance(result["skills"], list)
    assert isinstance(result["skills_grouped"], dict)
    assert set(result["skills_grouped"].keys()) == {"technical", "domain", "soft"}
    assert isinstance(result["sections"], dict)
    assert isinstance(result["extraction_warnings"], list)
    assert isinstance(result["readable"], bool)
    # Hard rule: never inferred from resume text - always None from the parser
    assert result["preferred_domain"] is None
    assert result["preferred_location"] is None


def test_pdf_response_matches_contract():
    pdf_bytes = create_synthetic_pdf(SAMPLE_TEXT_1)
    result = parse_resume(pdf_bytes, "pdf")
    assert_valid_contract(result)
    assert result["source_format"] == "pdf"


def test_docx_response_matches_contract():
    docx_bytes = create_synthetic_docx(SAMPLE_TEXT_1)
    result = parse_resume(docx_bytes, "docx")
    assert_valid_contract(result)
    assert result["source_format"] == "docx"


def test_scanned_pdf_response_matches_contract():
    pdf_bytes = create_synthetic_pdf("   ")
    result = parse_resume(pdf_bytes, "pdf")
    assert_valid_contract(result)
    assert result["readable"] is False
    assert result["skills"] == []
    assert result["skills_grouped"] == {"technical": [], "domain": [], "soft": []}
    assert "no_text_found_possibly_scanned" in result["extraction_warnings"]


def test_unsupported_file_type_matches_contract():
    result = parse_resume(b"irrelevant bytes", "txt")
    assert_valid_contract(result)
    assert result["readable"] is False
    assert result["skills"] == []
    assert "unsupported_file_type" in result["extraction_warnings"]


# ---------------------------------------------------------------------------
# Behavioral tests: does the parser actually extract the right content
# ---------------------------------------------------------------------------

def test_pdf_extracts_expected_sections():
    pdf_bytes = create_synthetic_pdf(SAMPLE_TEXT_1)
    result = parse_resume(pdf_bytes, "pdf")
    assert result["sections"].get("experience") is not None
    assert "Software Engineer" in result["sections"]["experience"]
    assert result["sections"].get("education") is not None
    assert "University of Tech" in result["sections"]["education"]


def test_docx_extracts_expected_sections():
    docx_bytes = create_synthetic_docx(SAMPLE_TEXT_1)
    result = parse_resume(docx_bytes, "docx")
    assert result["sections"].get("experience") is not None
    assert "docx_fallback_used" in result["extraction_warnings"]


def test_table_based_docx_still_produces_output():
    """Table-based resumes are a known harder case (content spread across
    cells rather than plain paragraphs) - this just confirms it doesn't
    crash and produces SOME readable content, not exact section accuracy."""
    docx_bytes = create_synthetic_docx(SAMPLE_TEXT_1, use_tables=True)
    result = parse_resume(docx_bytes, "docx")
    assert_valid_contract(result)
    assert result["raw_text"] != ""


def test_low_confidence_segmentation_flagged():
    """A document with no real section headers should get flagged, not
    silently mis-segmented."""
    text = "Just some random text\nwith no headings\nthat is very short."
    pdf_bytes = create_synthetic_pdf(text)
    result = parse_resume(pdf_bytes, "pdf")
    assert "low_section_match_confidence" in result["extraction_warnings"]


# ---------------------------------------------------------------------------
# Skill extraction tests: the actual integration with extract_skills.py
# ---------------------------------------------------------------------------

def test_known_skills_are_detected_from_resume():
    """The core integration point: skills genuinely present in the resume
    text, and present in master_skills.csv, must show up in the final
    skills list. This is what actually connects the resume parser to the
    rest of the pipeline (recommendation engine, role classifier)."""
    pdf_bytes = create_synthetic_pdf(SAMPLE_TEXT_1)
    result = parse_resume(pdf_bytes, "pdf")
    assert "Python" in result["skills"]
    assert "SQL" in result["skills"]


def test_skills_are_grouped_by_category():
    pdf_bytes = create_synthetic_pdf(SAMPLE_TEXT_1)
    result = parse_resume(pdf_bytes, "pdf")
    technical_names = [s["name"] for s in result["skills_grouped"]["technical"]]
    assert "Python" in technical_names or "SQL" in technical_names


def test_extract_resume_skills_empty_input():
    """Calling the skill-extraction function directly with empty text
    should return empty results, not error."""
    skills, grouped = extract_resume_skills("")
    assert skills == []
    assert grouped == {"technical": [], "domain": [], "soft": []}


def test_extract_resume_skills_direct_call():
    """Direct unit test of the skill-extraction function, independent of
    PDF/DOCX parsing - isolates whether the shared matcher integration
    itself works correctly."""
    skills, grouped = extract_resume_skills("Experienced with Python, Docker, and Machine Learning.")
    assert "Python" in skills
    assert "Docker" in skills
    assert "Machine Learning" in skills